"""
"과제 직무/대상자 검증" 탭 핵심 로직 — 업로드된 직무기술서(.docx 또는 .pdf)와
실제 인사데이터(과제 배정 인원 + 보유 전문성 분석)를 대조한다.

흐름:
  1) 업로드된 파일에서 텍스트를 뽑는다(extract_document) — 확장자로 분기한다.
     .docx는 표와 서술형 본문을 따로 뽑는다(_extract_docx). 표에는 보통
     "직무명 | 현인원 | 설명" 같은 구조화된 정보가 있고, 서술형 본문에는 더
     풍부한 역할 설명이 있는 경우가 많아 둘 다 쓴다. .pdf는 표/서술형 구분이
     없는 일반 텍스트 추출이라(_extract_pdf) 전부 서술형 본문으로 취급한다.
     지원하지 않는 확장자나 손상된 파일은 DocumentReadError로 화면에 바로
     보여줘도 안전한 한국어 메시지를 던진다.
  2) 표/서술형 각각을 LLM으로 구조화 추출(_extract_roles) — 직접 판단하지
     않고 문서에 적힌 내용만 그대로 뽑아내는 역할만 시킨다(이 프로젝트
     전체가 지켜온 "구조화 출력 강제" 원칙).
  3) 표와 서술형 추출 결과를 role_name 기준으로 병합(merge_roles) — 인원수는
     표를 authoritative로 삼고(사용자 확정), 표/서술형 인원수가 다르면
     "문서 내부 불일치"로 기록해 리포트에 노출한다.
  4) 선택한 과제(project_name — project_confl_address.csv 기준, 컨플루언스
     과제명과 동일한 체계)에 실제 배정된 인원을 researchers.csv의
     org_code로 조회한다(get_project_members) — process_project_expertise.py가
     문서 내 인력을 researcher_id로 매핑할 때 쓰는 것과 같은 기준이라, "이
     과제 사람만" 매칭 대상으로 한정된다(다른 과제 사람은 애초에 후보에
     안 들어옴). 결정적 로직, LLM 미사용.
  5) 이 과제의 컨플루언스 분석(project_expertise_analysis.json, 있으면)을
     읽어(_read_confluence_summary) 핵심기술/배경/산출물/난제/기대효과를
     맥락으로 삼아, 3)에서 병합한 직무 설명을 인사담당자가 이해하기 쉬운
     말로 다시 쓴다(_plain_explain_roles) — 원문(기술적일 수 있음)은
     raw_description으로 남겨 두고 대조해 볼 수 있게 한다.
  6) 각 실제 배정자마다 "가장 가까운 직무 1개"만 LLM이 판정한다
     (match_members_to_roles) — 여러 직무에 걸쳐 있어도 반드시 하나만
     고르고, 뚜렷한 근거가 없으면 "미분류"로 남긴다. 근거 문장도 쉬운
     말로 쓰도록 프롬프트에 명시했다. 이 매칭 인원수는 문서의 현인원과
     무관하게 독립적으로 계산된다 — 문서가 3명이라고 적었어도 실제
     매칭은 그보다 많거나 적거나 0명일 수 있고, 그 차이 자체가 리포트의
     핵심 정보다.
  7) 문서 수치 vs 데이터 기반 판정 수치를 대조한 리포트를 만들고
     (build_report) data/processed/jd_reconciliation/에 실행 이력으로 남긴다
     (save_history/list_history/load_history).

Source:
  data/processed/project_confl_address.csv     (과제 목록 — 드롭다운, dep_name/project_name)
  data/processed/researchers.csv               (과제 배정 인원 — org_code 기준, 이름/부서)
  data/processed/project_expertise_analysis.json (컨플루언스 과제 요약 — process_project_expertise.py)
  data/processed/연구원 보유 전문성 분석.json   (사람별 전문성 프로필 텍스트)
"""

import io
import json
import os
import re
import sys
from datetime import datetime

_PIPELINE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'pipeline')
sys.path.insert(0, os.path.abspath(_PIPELINE_DIR))

import llm_client  # noqa: E402
import pdf_reader  # noqa: E402
import researcher_fit as fit  # noqa: E402
from paths import OUT_DIR  # noqa: E402
from services import data_store  # noqa: E402

HISTORY_DIR = os.path.join(data_store.DATA_DIR, 'jd_reconciliation')


# ─── 과제/인원 조회 (결정적, LLM 미사용) ─────────────────────────────────────

def list_project_names() -> list:
    """드롭다운에 보여줄 과제 목록 — project_confl_address.csv(컨플루언스
    과제 목록, process_project_expertise.py가 분석하는 것과 동일한 과제명
    체계)의 project_name. 파일이 없으면(process_project_confl.py 미실행)
    빈 리스트."""
    projects_df = data_store.read_processed('project_confl_address')
    if projects_df.empty or 'project_name' not in projects_df.columns:
        return []
    return sorted({str(v).strip() for v in projects_df['project_name'] if str(v).strip()})


def get_project_members(project_name: str) -> list:
    """선택한 과제에 현재 배정된 인원 목록 — researchers.csv의
    org_code(현재 소속 과제)가 이 과제명과 같은 사람. process_project_expertise.py가
    문서 내 인력 이름을 researcher_id로 매핑할 때 쓰는 것과 동일한 기준이라,
    "이 과제 소속 사람만" 매칭 후보로 한정된다(다른 과제 사람은 애초에
    이 목록에 없음).

    project_name(project_confl_address.csv, "[탐색] 가나다라마바사"처럼 대괄호
    태그+띄어쓰기가 있을 수 있음)과 org_code(researchers.csv, "가나다라마바사"처럼
    태그/공백 없음)는 표기 형식이 달라 fit.normalize_org_code()로 정규화한
    뒤 비교한다(process_project_expertise.py의 _resolve_personnel()과 동일한
    정규화 기준 공유)."""
    researchers_df = data_store.read_processed('researchers')
    if researchers_df.empty:
        return []
    target = fit.normalize_org_code(project_name)
    rows = researchers_df[researchers_df['org_code'].astype(str).apply(fit.normalize_org_code) == target]
    if rows.empty:
        return []
    return [
        {
            'researcher_id': row['researcher_id'],
            'name': row.get('name', row['researcher_id']),
            'department': row.get('department', ''),
        }
        for _, row in rows.iterrows()
    ]


def _read_confluence_summary(project_name: str) -> dict | None:
    """process_project_expertise.py가 만든 project_expertise_analysis.json에서
    이 과제의 컨플루언스 분석 항목을 찾는다. 파일이 없거나(파이프라인 미실행)
    이 과제명이 없으면 None — 호출부는 이를 "컨플루언스 맥락 없이 직무기술서
    내용만으로 설명"으로 처리한다."""
    path = os.path.join(OUT_DIR, 'project_expertise_analysis.json')
    if not os.path.isfile(path):
        return None
    try:
        with open(path, encoding='utf-8') as f:
            items = json.load(f)
    except (OSError, json.JSONDecodeError):
        return None
    target = str(project_name or '').strip()
    for it in items:
        if str(it.get('project_name') or '').strip() == target:
            return it
    return None


def _confluence_context_text(entry: dict | None) -> str:
    """컨플루언스 과제 요약 항목을 LLM 맥락용 텍스트로 정리한다. "확인 불가"/
    빈 값은 제외한다(process_project_expertise.py가 추출 실패 시 채워 넣는
    placeholder이므로 맥락에 넣을 가치가 없음)."""
    if not entry:
        return ''

    def _ok(v):
        v = str(v or '').strip()
        return v if v and v != '확인 불가' else ''

    lines = []
    if _ok(entry.get('core_tech')):
        lines.append(f"핵심 기술: {_ok(entry.get('core_tech'))}")
    if _ok(entry.get('background')):
        lines.append(f"배경: {_ok(entry.get('background'))}")
    if _ok(entry.get('deliverable')):
        lines.append(f"최종 산출물: {_ok(entry.get('deliverable'))}")
    if _ok(entry.get('challenge')):
        lines.append(f"기술적 난제: {_ok(entry.get('challenge'))}")
    if _ok(entry.get('expected_impact')):
        lines.append(f"기대 효과: {_ok(entry.get('expected_impact'))}")
    milestones = [str(m).strip() for m in (entry.get('milestones') or []) if str(m).strip()]
    if milestones:
        lines.append('주요 마일스톤: ' + '; '.join(milestones))
    return '\n'.join(lines)


# ─── 문서 추출 (.docx / .pdf) ────────────────────────────────────────────────

class DocumentReadError(RuntimeError):
    """업로드된 파일에서 텍스트를 추출하지 못했을 때 — 메시지를 화면에 그대로
    노출해도 되는 사용자 대상 한국어 문구로 만든다(스택 트레이스 대신)."""


_SUPPORTED_EXTENSIONS = ('.docx', '.pdf')


def _extract_docx(file_bytes: bytes) -> dict:
    """표와 서술형 본문을 각각 텍스트로 뽑는다. 표는 셀을 ' | '로 이어 붙여
    LLM이 열 구조를 스스로 해석하게 한다(헤더 이름/열 순서가 문서마다 달라
    고정 파싱은 깨지기 쉬움)."""
    import zipfile

    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ModuleNotFoundError as exc:
        raise DocumentReadError(
            'python-docx 패키지가 설치되어 있지 않아 .docx 파일을 읽을 수 없습니다. '
            '서버에서 pip install -r requirements.txt(또는 pip install python-docx)를 '
            '실행한 뒤 앱을 재시작해주세요.'
        ) from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        # .docx는 내부적으로 zip 압축 포맷(OOXML)이다 — 확장자만 .docx로 바뀐
        # 옛 .doc(바이너리)이거나 파일이 손상되면 zip으로도 못 열려
        # BadZipFile이, zip은 맞지만 OOXML 필수 구성요소가 없으면
        # PackageNotFoundError가 난다. 둘 다 사용자 입장에선 같은 문제.
        raise DocumentReadError(
            '올바른 .docx 파일이 아닌 것 같습니다(예: 오래된 .doc 형식이거나 파일이 '
            '손상됨). Word에서 "다른 이름으로 저장" → .docx 형식으로 다시 저장한 뒤 '
            '업로드해주세요.'
        ) from exc
    except Exception as exc:  # noqa: BLE001 — 업로드 파일 파싱은 외부 입력 경계라 방어적으로 처리
        raise DocumentReadError(f'.docx 파일을 읽는 중 오류가 발생했습니다: {exc}') from exc

    narrative_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_blocks = []
    for table in doc.tables:
        rows_text = [' | '.join(c.text.strip() for c in row.cells) for row in table.rows]
        rows_text = [r for r in rows_text if r.strip(' |')]
        if rows_text:
            table_blocks.append('\n'.join(rows_text))
    return {
        'narrative_text': '\n'.join(narrative_lines),
        'tables_text': '\n\n'.join(table_blocks),
    }


def _extract_pdf(file_bytes: bytes) -> dict:
    """.pdf는 표/서술형을 구분해서 뽑을 수 없으므로(일반 PDF에는 docx 같은
    구조화된 표 정보가 없음) 페이지 텍스트 전체를 서술형 본문으로 취급한다."""
    try:
        text = pdf_reader.extract_text_from_bytes(file_bytes, label='업로드한 PDF')
    except pdf_reader.PdfNotFoundError as exc:
        raise DocumentReadError(str(exc)) from exc
    return {'narrative_text': text, 'tables_text': ''}


def extract_document(file_bytes: bytes, filename: str) -> dict:
    """파일 확장자로 .docx/.pdf 추출기를 분기 호출한다. 지원하지 않는 확장자면
    DocumentReadError(화면에 바로 보여줘도 되는 메시지)."""
    ext = os.path.splitext(filename or '')[1].lower()
    if ext == '.docx':
        return _extract_docx(file_bytes)
    if ext == '.pdf':
        return _extract_pdf(file_bytes)
    raise DocumentReadError(
        f'지원하지 않는 파일 형식입니다({ext or "확장자 없음"}). '
        f'{"/".join(_SUPPORTED_EXTENSIONS)} 파일을 업로드해주세요.'
    )


# ─── ① 문서 → 직무별 구조화 추출 (LLM, 판단 없이 추출만) ─────────────────────

_ROLE_EXTRACTION_SYSTEM_PROMPT = """# Role
당신은 R&D 과제의 직무기술서(job description) 문서에서 직무(역할)별 필요/
현재 인원수와 역할 설명을 뽑아내는 "문서 구조화 도우미"입니다. 직접
판단하거나 추정하지 않고, 문서에 실제로 적혀 있는 내용만 그대로
추출합니다.

# Guidelines & Constraints
1. 문서에 없는 직무를 만들어내거나, 적혀 있지 않은 인원수를 추정하지 마세요.
2. 인원수가 명시되지 않은 직무는 headcount를 0으로 두세요.
3. description에는 그 직무의 역할/요구 전문성 설명을 문서 원문에 가깝게
   요약해 담으세요(다음 단계에서 이 설명으로 실제 인력과 매칭에 씁니다).
4. 반드시 아래 JSON 형식으로만 출력하고, 그 외 텍스트는 출력하지 마세요.

# Output Format (JSON)
{"roles": [{"role_name": "직무명", "headcount": 0, "description": "역할/요구 전문성 설명"}]}
"""


def _extract_roles(text: str, source_label: str) -> list:
    text = (text or '').strip()
    if not text:
        return []
    prompt = (
        f'[{source_label}에서 추출한 문서 내용]\n{text}\n\n'
        '위 내용에서 직무(역할)별 인원수와 역할 설명을 뽑아주세요.'
    )
    raw = llm_client.call_llm(prompt, _ROLE_EXTRACTION_SYSTEM_PROMPT, temperature=0.0, max_tokens=1500)
    if not raw:
        return []
    try:
        parsed = json.loads(llm_client.extract_json(raw))
    except json.JSONDecodeError:
        return []

    roles = []
    for r in parsed.get('roles') or []:
        name = str(r.get('role_name') or '').strip()
        if not name:
            continue
        try:
            headcount = int(r.get('headcount') or 0)
        except (TypeError, ValueError):
            headcount = 0
        roles.append({
            'role_name': name, 'headcount': headcount,
            'description': str(r.get('description') or '').strip(),
        })
    return roles


# ─── ③ 표/서술형 병합 + 문서 내부 정합성 체크 ────────────────────────────────

def _norm_role(name: str) -> str:
    return re.sub(r'\s+', '', str(name or '')).strip().lower()


def _pair_roles(table_roles: list, narrative_roles: list) -> list:
    """role_name 기준으로 표 항목과 서술형 항목을 짝짓는다(정규화 후 부분
    일치). 표에 없는 role_name이 서술형에만 있으면 (None, narrative_role)로
    남긴다."""
    pairs = []
    used = set()
    for t in table_roles:
        tn = _norm_role(t['role_name'])
        match_i = None
        for i, n in enumerate(narrative_roles):
            if i in used:
                continue
            nn = _norm_role(n['role_name'])
            if tn == nn or tn in nn or nn in tn:
                match_i = i
                break
        if match_i is not None:
            used.add(match_i)
            pairs.append((t, narrative_roles[match_i]))
        else:
            pairs.append((t, None))
    for i, n in enumerate(narrative_roles):
        if i not in used:
            pairs.append((None, n))
    return pairs


def merge_roles(table_roles: list, narrative_roles: list) -> tuple:
    """표의 headcount를 authoritative로 삼아 최종 직무 목록을 만들고, 표/
    서술형 인원수가 다르면 문서 내부 불일치 메모를 남긴다(사용자 확정:
    표를 우선하되 서술형과도 비교)."""
    pairs = _pair_roles(table_roles, narrative_roles)
    merged, notes = [], []
    for t, n in pairs:
        if t and n:
            role_name, headcount = t['role_name'], t['headcount']
            desc = '\n'.join(filter(None, [t['description'], n['description']]))
            if n['headcount'] and n['headcount'] != t['headcount']:
                notes.append(
                    f"'{role_name}' 인원수가 표에는 {t['headcount']}명, 본문 서술에는 "
                    f"{n['headcount']}명으로 다르게 적혀 있습니다(표 기준으로 비교합니다)."
                )
        elif t:
            role_name, headcount, desc = t['role_name'], t['headcount'], t['description']
        else:
            role_name, headcount, desc = n['role_name'], n['headcount'], n['description']
            notes.append(f"'{role_name}'은 표에는 없고 본문 서술에만 언급되어 있습니다.")
        merged.append({'role_name': role_name, 'headcount': headcount, 'description': desc})
    return merged, notes


# ─── ④ 컨플루언스 맥락 기반 쉬운 말 직무 설명 재작성 ─────────────────────────

_PLAIN_EXPLAIN_SYSTEM_PROMPT = """# Role
당신은 R&D 과제의 기술적인 직무 설명을, 기술 전문가가 아닌 HR 담당자도
바로 이해할 수 있도록 쉬운 일상 언어로 다시 설명해주는 "쉬운 설명
도우미"입니다.

# Goal
직무기술서에서 뽑은 직무별 설명(전문 용어가 섞여 있을 수 있음)을, 함께
주어지는 과제 배경 정보를 참고해 더 이해하기 쉬운 설명으로 다시
씁니다.

# Guidelines & Constraints
1. 원문에 없는 내용을 새로 지어내지 마세요 — 같은 의미를 쉬운 말로
   바꿔 쓰는 것이지, 추정하거나 과장하지 않습니다.
2. AI·기술 전문 용어(모델/알고리즘 이름, 프레임워크명 등)가 나오면
   "그게 무엇을 하는 것인지"를 함께 풀어서 설명하세요.
3. 과제 배경 정보는 문맥 이해를 돕는 참고용입니다 — 직무 설명 자체는
   원문 직무기술서 내용을 벗어나지 않아야 합니다.
4. 각 직무마다 2~4문장 정도로 정리하세요.
5. 반드시 아래 JSON 형식으로만 출력하고, 입력받은 모든 직무를 빠짐없이
   포함하세요.

# Output Format (JSON)
{"roles": [{"role_name": "직무명(입력과 동일하게)", "plain_description": "쉬운 말 설명"}]}
"""


def _plain_explain_roles(roles: list, confluence_context: str) -> list:
    """merge_roles()가 만든 직무 목록의 description을 쉬운 말로 다시 쓴다.
    원문은 raw_description으로 그대로 보존해, 화면에서 대조해 볼 수 있게
    한다. LLM 실패/미매칭 시 해당 직무는 원문 그대로 둔다(안전한 폴백)."""
    if not roles:
        return []

    role_block = '\n\n'.join(
        f'[{r["role_name"]}]\n{r["description"] or "(설명 없음)"}' for r in roles
    )
    context_block = confluence_context or '(이 과제의 컨플루언스 분석 데이터 없음)'
    prompt = (
        f'[과제 배경 정보(참고용)]\n{context_block}\n\n'
        f'[직무기술서에서 뽑은 직무별 설명(원문)]\n{role_block}\n\n'
        '위 직무들을 각각 쉬운 말로 다시 설명해주세요.'
    )
    raw = llm_client.call_llm(prompt, _PLAIN_EXPLAIN_SYSTEM_PROMPT, temperature=0.0, max_tokens=2000)
    plain_by_role = {}
    if raw:
        try:
            parsed = json.loads(llm_client.extract_json(raw))
            plain_by_role = {
                _norm_role(r.get('role_name')): str(r.get('plain_description') or '').strip()
                for r in (parsed.get('roles') or [])
            }
        except json.JSONDecodeError:
            plain_by_role = {}

    result = []
    for r in roles:
        plain = plain_by_role.get(_norm_role(r['role_name']))
        result.append({
            **r,
            'raw_description': r['description'],
            'description': plain or r['description'],
        })
    return result


# ─── ⑤ 사람 × 직무 매핑 (가까운 직무 1개만, 쉬운 말 근거) ─────────────────────

_MATCH_SYSTEM_PROMPT = """# Role
당신은 R&D 인력 배치 근거를 HR 담당자에게 설명하는 "R&D 직무 매칭
도우미"입니다. HR 담당자는 기술 전문가가 아니므로, 전문 용어(모델/알고리즘
이름 등 AI·기술 전문 용어)를 그대로 쓰지 말고 누구나 바로 이해할 수 있는
쉬운 일상 언어로 설명합니다.

# Goal
한 연구원의 보유 전문성 프로필과, 문서에 정의된 여러 직무(역할) 후보를
비교해 이 연구원에게 "가장 가까운 직무 딱 하나"만 고릅니다.

# Guidelines & Constraints
1. 철저한 팩트 기반: 프로필에 없는 역량을 있다고 추정하지 마세요.
2. 여러 직무에 걸쳐 있어 보여도 반드시 하나만 선택하세요(가장 가까운 것).
3. 근거가 뚜렷하면 confidence를 "상", 관련은 있지만 확신이 낮으면 "중",
   거의 근거가 없으면 "하"로 낮추고 best_candidate를 "미분류"로 답하세요.
4. evidence는 전문 용어 없이 1~2문장의 쉬운 말로 쓰세요
   (예: "논문 검증·재현 경험이 많아 결과를 확인하는 역할에 가깝습니다").
5. 반드시 아래 JSON 형식으로만 출력하세요.

# Output Format (JSON)
{"best_candidate": "후보 라벨 또는 미분류", "confidence": "상|중|하", "evidence": "쉬운 말 설명"}
"""


def _label_of(i: int) -> str:
    return chr(65 + i) if i < 26 else f'X{i}'


def match_members_to_roles(members: list, roles: list) -> list:
    """반환: [{'researcher_id','role_name'(또는 '미분류'),'confidence','evidence'}, ...]"""
    if not members:
        return []
    if not roles:
        return [
            {'researcher_id': m['researcher_id'], 'role_name': '미분류',
             'confidence': '하', 'evidence': '문서에서 인식된 직무가 없습니다.'}
            for m in members
        ]

    labels = [f'후보 {_label_of(i)}' for i in range(len(roles))]
    role_block = '\n\n'.join(
        f'[{lbl}] {r["role_name"]}\n{r["description"] or "(설명 없음)"}'
        for lbl, r in zip(labels, roles)
    )
    label_to_role = dict(zip(labels, [r['role_name'] for r in roles]))
    profiles = data_store.read_expertise_profiles()

    def _judge(profile_text: str):
        prompt = (
            f'[연구원 전문성 프로필]\n{profile_text}\n\n[후보 직무 목록]\n{role_block}\n\n'
            '위 연구원에게 가장 가까운 직무 하나를 골라주세요.'
        )
        raw = llm_client.call_llm(prompt, _MATCH_SYSTEM_PROMPT, temperature=0.0, max_tokens=600)
        if not raw:
            return None
        try:
            return json.loads(llm_client.extract_json(raw))
        except json.JSONDecodeError:
            return None

    tasks = []
    for m in members:
        profile = profiles.get(m['researcher_id'])
        tasks.append((lambda pt=fit.researcher_profile_text(profile): _judge(pt)) if profile else None)

    real_tasks = [t for t in tasks if t is not None]
    workers = llm_client.max_concurrency()
    task_results = llm_client.run_concurrent(real_tasks, max_workers=workers) if real_tasks else []

    results, it = [], iter(task_results)
    for m, t in zip(members, tasks):
        if t is None:
            results.append({'researcher_id': m['researcher_id'], 'role_name': '미분류',
                             'confidence': '하', 'evidence': '전문성 분석 데이터가 없습니다.'})
            continue
        parsed, error = next(it)
        if error is not None or not parsed:
            results.append({'researcher_id': m['researcher_id'], 'role_name': '미분류',
                             'confidence': '하', 'evidence': '판정 중 오류가 발생했습니다.'})
            continue
        best = str(parsed.get('best_candidate') or '').strip()
        results.append({
            'researcher_id': m['researcher_id'],
            'role_name': label_to_role.get(best, '미분류'),
            'confidence': parsed.get('confidence', ''),
            'evidence': parsed.get('evidence', ''),
        })
    return results


# ─── ⑥ 리포트 조립 + 이력 저장/조회 ─────────────────────────────────────────

def build_report(project_name: str, doc_filename: str, table_roles: list, narrative_roles: list,
                  roles: list, consistency_notes: list, members: list, matches: list,
                  confluence_entry: dict | None = None) -> dict:
    member_by_id = {m['researcher_id']: m for m in members}

    role_rows = []
    for r in roles:
        matched = [
            {**member_by_id[mt['researcher_id']], 'confidence': mt['confidence'], 'evidence': mt['evidence']}
            for mt in matches if mt['role_name'] == r['role_name']
        ]
        role_rows.append({
            'role_name': r['role_name'],
            'description': r['description'],
            'raw_description': r.get('raw_description', r['description']),
            'document_count': r['headcount'],
            'matched_count': len(matched),
            'diff': len(matched) - r['headcount'],
            'matched': matched,
        })
    unmatched = [
        {**member_by_id[mt['researcher_id']], 'confidence': mt['confidence'], 'evidence': mt['evidence']}
        for mt in matches if mt['role_name'] == '미분류' and mt['researcher_id'] in member_by_id
    ]

    document_total = sum(r['headcount'] for r in roles)
    actual_total = len(members)
    summary_text = (
        f'문서상 이 과제는 총 {document_total}명이 필요하다고 되어 있고, '
        f'실제로는 {actual_total}명이 배정되어 있습니다.'
    )
    if unmatched:
        summary_text += f' 그중 {len(unmatched)}명은 문서에 적힌 어떤 직무와도 뚜렷이 맞지 않는 것으로 판단됩니다.'
    if consistency_notes:
        summary_text += ' 문서 내부에 표와 본문 내용이 다른 부분이 있어 함께 확인이 필요합니다.'
    if not confluence_entry:
        summary_text += ' 이 과제의 컨플루언스 분석 데이터가 없어 직무기술서 내용만으로 설명했습니다.'

    return {
        'project_name': project_name,
        'doc_filename': doc_filename,
        'run_at': datetime.now().isoformat(timespec='seconds'),
        'table_roles': table_roles,
        'narrative_roles': narrative_roles,
        'consistency_notes': consistency_notes,
        'roles': role_rows,
        'unmatched': unmatched,
        'document_total': document_total,
        'actual_total': actual_total,
        'summary_text': summary_text,
        'confluence_available': bool(confluence_entry),
        'project_overview': _confluence_context_text(confluence_entry),
    }


def _safe_stem(text: str) -> str:
    return re.sub(r'[^0-9A-Za-z가-힣_-]', '_', text or '')[:60] or 'project'


def save_history(report: dict) -> str:
    os.makedirs(HISTORY_DIR, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    fname = f"{_safe_stem(report.get('project_name', ''))}_{ts}.json"
    with open(os.path.join(HISTORY_DIR, fname), 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    return fname


def list_history() -> list:
    if not os.path.isdir(HISTORY_DIR):
        return []
    rows = []
    for fname in os.listdir(HISTORY_DIR):
        if not fname.endswith('.json'):
            continue
        try:
            with open(os.path.join(HISTORY_DIR, fname), encoding='utf-8') as f:
                data = json.load(f)
        except (OSError, json.JSONDecodeError):
            continue
        rows.append({
            'file': fname,
            'run_at': data.get('run_at', ''),
            'project_name': data.get('project_name', ''),
            'doc_filename': data.get('doc_filename', ''),
            'document_total': data.get('document_total', 0),
            'actual_total': data.get('actual_total', 0),
        })
    rows.sort(key=lambda r: r['run_at'], reverse=True)
    return rows


def load_history(file_name: str) -> dict | None:
    """file_name은 list_history()가 돌려준 값만 받는다 — basename만 취급하고
    HISTORY_DIR 밖 경로는 거부해 경로 조작을 막는다."""
    safe_name = os.path.basename(file_name or '')
    if not safe_name.endswith('.json'):
        return None
    path = os.path.join(HISTORY_DIR, safe_name)
    if not os.path.isfile(path):
        return None
    with open(path, encoding='utf-8') as f:
        return json.load(f)


# ─── 전체 실행 ───────────────────────────────────────────────────────────────

def run_reconciliation(project_name: str, doc_filename: str, file_bytes: bytes) -> dict:
    members = get_project_members(project_name)
    extracted = extract_document(file_bytes, doc_filename)
    table_roles = _extract_roles(extracted['tables_text'], '표')
    narrative_roles = _extract_roles(extracted['narrative_text'], '서술형 본문')
    roles, consistency_notes = merge_roles(table_roles, narrative_roles)
    confluence_entry = _read_confluence_summary(project_name)
    roles = _plain_explain_roles(roles, _confluence_context_text(confluence_entry))
    matches = match_members_to_roles(members, roles)
    report = build_report(project_name, doc_filename, table_roles, narrative_roles,
                           roles, consistency_notes, members, matches, confluence_entry)
    save_history(report)
    return report
